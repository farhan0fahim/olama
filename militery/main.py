import os, requests, uvicorn, time, sys, threading, re
from fastapi import FastAPI, Request, Form, Query
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.templating import Jinja2Templates
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import List
from bs4 import BeautifulSoup
from urllib.parse import urlparse
from apscheduler.schedulers.background import BackgroundScheduler
from transformers import pipeline

# Windows Unicode Terminal Support
if sys.platform == "win32":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

app = FastAPI()
templates = Jinja2Templates(directory="templates")

# --- SYSTEM CORE STATE ---
summarizer = None
is_model_ready = False
SYSTEM_LOGS = []
INTEL_CACHE = [] 
SYNC_INTERVAL_MINS = 5 
trigger_sync_now = threading.Event()

# Initial Deploy Grid
TARGET_GRID = {
    "Prothom Alo": {
        "type": "national",
        "paths": {"Politics": "https://www.prothomalo.com/politics", "Economy": "https://www.prothomalo.com/business", "National": "https://www.prothomalo.com/bangladesh"}
    }
}

def add_log(msg):
    global SYSTEM_LOGS
    t = datetime.now().strftime("%H:%M:%S")
    entry = f"[{t}] {msg}"
    print(entry)
    SYSTEM_LOGS.append(entry)
    if len(SYSTEM_LOGS) > 30: SYSTEM_LOGS.pop(0)

# --- AI SYNTHESIS KERNEL ---
def load_ai():
    global summarizer, is_model_ready
    add_log("AI_KERNEL: LOADING MULTILINGUAL NEWS SYNTHESIZER...")
    try:
        summarizer = pipeline("summarization", model="csebuetnlp/mT5_multilingual_XLSum", framework="pt", device=-1)
        is_model_ready = True
        add_log("AI_KERNEL: ONLINE. HIGH-PRECISION ANALYSIS ENABLED.")
    except Exception as e:
        add_log(f"AI_KERNEL_ERROR: {e}")

threading.Thread(target=load_ai, daemon=True).start()

# --- HEURISTIC DEEP SCRAPER (Universal News Discovery) ---
def universal_intercept(url, source_name):
    if not url: return []
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'}
    news = []
    
    # 1. API Mode (Prothom Alo optimization)
    if "prothomalo" in url:
        slug = urlparse(url).path.strip("/")
        if "business" in url or "economy" in url: slug = "business"
        try:
            api_url = f"https://www.prothomalo.com/api/v1/collections/{slug}?offset=0&limit=5"
            data = requests.get(api_url, timeout=5).json()
            for item in data.get('items', []):
                story = item.get('story', {})
                if story.get('headline'):
                    news.append({"title": story.get('headline'), "link": f"https://www.prothomalo.com/{story.get('slug')}"})
            if news: return news
        except: pass

    # 2. Heuristic Scrape Mode (Works for Al Jazeera, BBC, Reuters, etc.)
    try:
        res = requests.get(url, headers=headers, timeout=12, verify=True)
        soup = BeautifulSoup(res.text, 'html.parser')
        
        # Remove noise (scripts, styles, nav, footer)
        for noise in soup(['script', 'style', 'nav', 'footer', 'header']):
            noise.decompose()

        # Discovery Logic: Find links with high word count and structural importance
        links = soup.find_all('a', href=True)
        for a in links:
            title = a.get_text().strip()
            # Headline Heuristic: 
            # - Must be longer than 35 chars
            # - Must contain at least 4 spaces (multiple words)
            # - Must not be a repetition
            if len(title) > 35 and title.count(' ') >= 4:
                href = a['href']
                if not href.startswith('http'):
                    p = urlparse(url)
                    href = f"{p.scheme}://{p.netloc}" + (href if href.startswith('/') else '/' + href)
                
                if not any(n['link'] == href for n in news):
                    news.append({"title": title, "link": href})
            
            if len(news) >= 6: break # Collect top 6 relevant packets
            
    except Exception as e:
        add_log(f"INTERCEPT_FAILURE: {source_name} unreachable.")
        
    return news

def get_deep_summary(url):
    if not is_model_ready: return "AI CORE INITIALIZING. PREVIEW UNAVAILABLE."
    try:
        res = requests.get(url, headers={'User-Agent': 'Mozilla/5.0'}, timeout=8)
        soup = BeautifulSoup(res.text, 'html.parser')
        paragraphs = soup.find_all('p')
        # Analyze first 8 paragraphs for a truly deep 5-10 line summary
        txt = " ".join([p.text for p in paragraphs[:8]]) 
        if len(txt) < 200: return "Dossier Notice: Link contains insufficient textual data for AI synthesis."
        
        # Enforcing 5-10 line output via max_length
        summary = summarizer(txt[:2500], max_length=350, min_length=150, length_penalty=2.5, do_sample=False)
        return summary[0]['summary_text']
    except: return "Deep Link Analysis Interrupted."

# --- SEQUENTIAL BACKGROUND CACHE ---
def background_worker():
    global INTEL_CACHE
    while True:
        temp_cache = []
        add_log("SYNC_ENGINE: PROBING GLOBAL TARGET GRID...")
        current_targets = dict(TARGET_GRID)
        
        for name, cfg in current_targets.items():
            for sector, url in cfg['paths'].items():
                add_log(f"INTERCEPTING: {name.upper()} -> {sector.upper()}")
                headlines = universal_intercept(url, name)
                
                for item in headlines[:3]:
                    add_log(f"SYNTHESIZING: {item['title'][:40]}...")
                    summary = get_deep_summary(item['link'])
                    temp_cache.append({
                        **item, "source": name, "sector": sector, 
                        "summary": summary, "type": cfg['type'],
                        "time": datetime.now().strftime("%H:%M")
                    })
        
        INTEL_CACHE = temp_cache
        add_log(f"SYNC_ENGINE: CACHE REFRESHED. {len(INTEL_CACHE)} PACKETS READY.")
        
        # Trigger-based wait
        trigger_sync_now.wait(timeout=SYNC_INTERVAL_MINS * 60)
        trigger_sync_now.clear()

threading.Thread(target=background_worker, daemon=True).start()

# --- REPORT AUTOMATION ---
def auto_archive():
    if not INTEL_CACHE: return
    add_log("ARCHIVER: AUTO-GENERATING PERIODIC DOSSIER...")
    doc = Document()
    doc.add_heading('PERIODIC INTELLIGENCE LOG', 0)
    for n in INTEL_CACHE:
        doc.add_heading(f"[{n['source']}] {n['title']}", level=3)
        doc.add_paragraph(n['summary'])
    fname = f"ARCHIVE_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    doc.save(fname)

scheduler = BackgroundScheduler()
scheduler.add_job(auto_archive, 'interval', minutes=15, id='archive_job')
scheduler.start()

# --- API ROUTES ---
@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request, "outlets": TARGET_GRID})

@app.get("/fetch-news")
async def fetch_news(selected_sources: List[str] = Query([]), selected_sectors: List[str] = Query([])):
    # Match cache against UI filters
    filtered = [n for n in INTEL_CACHE if n['source'] in selected_sources and (n['sector'] in selected_sectors or n['type'] == 'international')]
    return {"news": filtered, "logs": SYSTEM_LOGS}

@app.post("/update-sync-timer")
async def update_sync_timer(minutes: int = Form(...)):
    global SYNC_INTERVAL_MINS
    SYNC_INTERVAL_MINS = max(1, minutes)
    add_log(f"SYSTEM: SYNC INTERVAL SET TO {minutes} MINS.")
    return {"status": "success"}

@app.post("/force-sync")
async def force_sync():
    add_log("SYSTEM: MANUAL SYNC OVERRIDE ACTIVATED.")
    trigger_sync_now.set()
    return {"status": "syncing"}

@app.post("/update-archive-timer")
async def update_archive_timer(minutes: int = Form(...)):
    scheduler.reschedule_job('archive_job', trigger='interval', minutes=max(1, minutes))
    add_log(f"SYSTEM: ARCHIVE INTERVAL SET TO {minutes} MINS.")
    return {"status": "success"}

@app.post("/generate-intel")
async def generate_intel(sources: List[str] = Form(...), sectors: List[str] = Form(...)):
    add_log("DOSSIER: COMPILING STRUCTURED DOCX...")
    doc = Document()
    doc.add_heading('TOP SECRET // INTELLIGENCE BRIEFING', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Section 1: International
    doc.add_heading('SECTION I: INTERNATIONAL INTELLIGENCE', level=1)
    for n in INTEL_CACHE:
        if n['source'] in sources and n['type'] == 'international':
            doc.add_heading(n['title'], level=3)
            doc.add_paragraph(f"SOURCE: {n['source']} | LINK: {n['link']}", style='Caption')
            doc.add_paragraph(n['summary'])

    # Section 2: National
    doc.add_heading('SECTION II: NATIONAL INTELLIGENCE', level=1)
    for sec in ["Politics", "National", "Economy"]:
        if sec in sectors:
            doc.add_heading(f"SECTOR: {sec.upper()}", level=2)
            for n in INTEL_CACHE:
                if n['source'] in sources and n['sector'] == sec:
                    p = doc.add_paragraph()
                    p.add_run(f"[{n['source']}] {n['title']}").bold = True
                    doc.add_paragraph(n['summary'])

    fname = f"INTEL_DOSSIER_{datetime.now().strftime('%H%M')}.docx"
    doc.save(fname)
    return FileResponse(fname, filename=fname)

@app.post("/add-paper")
async def add_paper(name: str = Form(...), type: str = Form(...), pol: str = Form(""), eco: str = Form(""), nat: str = Form(""), intl_url: str = Form("")):
    if type == "national":
        TARGET_GRID[name] = {"type": "national", "paths": {"Politics": pol, "Economy": eco, "National": nat}}
    else:
        TARGET_GRID[name] = {"type": "international", "paths": {"International": intl_url}}
    return {"status": "success"}

@app.post("/delete-paper")
async def delete_paper(name: str = Form(...)):
    if name in TARGET_GRID: del TARGET_GRID[name]
    return {"status": "success"}

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)